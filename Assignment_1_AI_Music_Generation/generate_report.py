"""
Generate PHKI Assignment 1 Analytical Report as .docx
Elias Schwegler - HSLU FS26
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import os
import copy

# ── Paths ──────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "PHKI_Assignment1_Report.docx")

# ── Colours ────────────────────────────────────────────
NAVY = RGBColor(0x1F, 0x39, 0x64)
DARK_GREY = RGBColor(0x40, 0x40, 0x40)
BLACK = RGBColor(0x00, 0x00, 0x00)

doc = Document()

# ── Footnote Infrastructure ────────────────────────────
# python-docx has no native footnote API, so we build it via XML.

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

_footnote_id = [0]  # mutable counter


def _ensure_footnotes_part():
    """Create the footnotes.xml part if it doesn't exist yet."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    part_name = PackURI('/word/footnotes.xml')
    # Check if already exists
    for rel in doc.part.rels.values():
        try:
            if rel.target_part.partname == part_name:
                return rel.target_part
        except (AttributeError, KeyError):
            continue

    # Create minimal footnotes.xml with separator footnotes (id 0 and 1)
    xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        '<w:footnote w:type="separator" w:id="0">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="1">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )
    fn_part = Part(
        part_name,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
        xml.encode('utf-8'),
        doc.part.package,
    )
    doc.part.relate_to(fn_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes')
    return fn_part


def add_footnote(paragraph, text):
    """Add a footnote to the end of a paragraph with the given text."""
    fn_part = _ensure_footnotes_part()
    _footnote_id[0] += 2  # start at 2 (0 and 1 are separators)
    fn_id = _footnote_id[0]

    # 1. Add the footnote content to footnotes.xml
    footnotes_elem = etree.fromstring(fn_part.blob)
    fn_elem = etree.SubElement(footnotes_elem, qn('w:footnote'))
    fn_elem.set(qn('w:id'), str(fn_id))

    fn_p = etree.SubElement(fn_elem, qn('w:p'))
    # Footnote ref run
    fn_r1 = etree.SubElement(fn_p, qn('w:r'))
    fn_rpr1 = etree.SubElement(fn_r1, qn('w:rPr'))
    fn_style1 = etree.SubElement(fn_rpr1, qn('w:rStyle'))
    fn_style1.set(qn('w:val'), 'FootnoteReference')
    etree.SubElement(fn_r1, qn('w:footnoteRef'))
    # Space
    fn_r2 = etree.SubElement(fn_p, qn('w:r'))
    fn_t2 = etree.SubElement(fn_r2, qn('w:t'))
    fn_t2.text = ' ' + text
    fn_t2.set(qn('xml:space'), 'preserve')
    # Set footnote text size to 9pt
    fn_rpr2 = etree.SubElement(fn_r2, qn('w:rPr'))
    fn_sz = etree.SubElement(fn_rpr2, qn('w:sz'))
    fn_sz.set(qn('w:val'), '18')  # half-points

    fn_part._blob = etree.tostring(footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 2. Add footnote reference in the paragraph
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rstyle = OxmlElement('w:rStyle')
    rstyle.set(qn('w:val'), 'FootnoteReference')
    rpr.append(rstyle)
    run.append(rpr)
    fn_ref = OxmlElement('w:footnoteReference')
    fn_ref.set(qn('w:id'), str(fn_id))
    run.append(fn_ref)
    paragraph._p.append(run)

    return paragraph


# ── Page Setup ─────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ─────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = BLACK
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for level, size in [(1, 16), (2, 13), (3, 11)]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.size = Pt(size)
    hs.font.color.rgb = BLACK
    hs.font.bold = True
    if level == 3:
        hs.font.italic = True


def add_para(text, bold=False, italic=False, align=None, size=None, color=None, space_after=None):
    """Helper to add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_mixed_para(parts, align=None):
    """Add a paragraph with mixed formatting. parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    if align:
        p.alignment = align
    return p


# ══════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════

for _ in range(6):
    add_para("", space_after=0)

add_para(
    "Synthetic Pine & Heavy Air",
    bold=True, size=26, color=NAVY,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
)

add_para(
    "An Analytical Report on AI Music Generation with Google Gemini",
    italic=True, size=14, color=DARK_GREY,
    align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
)

for line in [
    "Elias Schwegler",
    "",
    "Philosophy, Art, and Artificial Intelligence (PHKI FS26)",
    "HSLU - Lucerne University of Applied Sciences and Arts",
    "April 2026",
    "",
    "Lecturers: Shaelom Fischer, Guillaume Massol, Catherine Hayden",
]:
    add_para(line, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════

doc.add_heading("1. Introduction", level=1)

p = doc.add_paragraph(
    "Some songs define a generation. Passenger\u2019s \u201cLet Her Go\u201d (Passenger, 2012) is one of them  - a "
    "melancholic folk-pop ballad that has accumulated nearly five billion views on YouTube and became the "
    "soundtrack to countless adolescent memories, including my own. When Google Gemini introduced its music "
    "generation capabilities, "
    "I saw an opportunity to ask a question that sits at the intersection of technology, art, and ethics: "
    "can an AI convincingly replicate the emotional signature of a chart-topping song? And if it can, what "
    "does that mean for the artists whose voices, styles, and creative labour made that replication possible?"
)
add_footnote(p, "Passenger (2012). Let Her Go [Song]. On All the Little Lights. Black Crow Records.")

doc.add_paragraph(
    "To find out, I designed a two-part experiment. First, I asked Gemini to create a song with the same "
    "melody as \u201cLet Her Go\u201d but with entirely different, deliberately absurd lyrics  - about a car, its "
    "tires, and a synthetic pine air freshener. Second, I reversed the test: I asked for a new melody set to "
    "the original copyrighted lyrics. Gemini refused both requests on copyright grounds, yet in each case "
    "produced an original track that captured the emotional atmosphere of the reference song with startling "
    "fidelity. This report analyses the creative process behind those experiments, the role Gemini played as "
    "an AI co-creator, and how the results connect to the philosophical and cultural themes explored "
    "throughout the PHKI module."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  2. ANALYSIS OF THE CREATIVE PROCESS
# ══════════════════════════════════════════════════════════

doc.add_heading("2. Analysis of the Creative Process", level=1)

doc.add_heading("2.1 Experimental Design", level=2)

doc.add_paragraph(
    "The experiment was deliberately designed as a stress test rather than an open-ended creative exploration. "
    "I chose \u201cLet Her Go\u201d by Passenger precisely because of its ubiquity  - a song so widely known that any "
    "resemblance in the AI\u2019s output would be immediately recognisable. The goal was not to create the most "
    "beautiful piece of AI-generated music, but to probe where Gemini\u2019s creative capabilities end and its "
    "copyright guardrails begin."
)

doc.add_heading("2.2 Test 1: Same Melody, Different Text", level=2)

doc.add_paragraph(
    "In the first test (see Appendix: Chat_1_01.png, Chat_1_02.png), I provided Gemini with a YouTube link "
    "to \u201cLet Her Go\u201d and asked it to create a song with the same melody but different, melancholic lyrics  - "
    "specifically about sitting in a car, the friction of the tires on the road, and the fading scent of a "
    "synthetic pine air freshener hanging from the rearview mirror. I chose this absurd subject matter "
    "intentionally: if the AI could make a song about a pine air freshener feel genuinely moving, that would "
    "demonstrate its ability to replicate emotional atmosphere independent of lyrical content."
)

doc.add_paragraph(
    "Gemini\u2019s response was revealing. It explicitly stated that it could not recreate the exact melody or "
    "use the song\u2019s intellectual property. Instead, it generated a brand-new track: \u201cSynthetic Pine & Heavy "
    "Air,\u201d a melancholic Bedroom Pop song at 80 BPM featuring warm, reverberant electric piano, delicate "
    "guitar textures with light vibrato, and breathy, close-miked vocals. The track builds from a sparse, "
    "atmospheric intro into a cinematic climax before fading out. The result was, frankly, unsettling in its "
    "quality. The emotional register  - quiet solitude, wistful reflection, the weight of things left unsaid "
    " - matched the reference song almost perfectly, even though the melody and lyrics were entirely original."
)

doc.add_heading("2.3 Test 2: New Melody, Same Lyrics", level=2)

doc.add_paragraph(
    "To test the boundaries from the opposite direction (see Appendix: Chat_1_03_follow_up.png, "
    "Chat_1_04_follow_up.png), I asked Gemini to create a new melody for the original \u201cLet Her Go\u201d lyrics. "
    "Again, Gemini refused, stating that it could not use the copyrighted lyrics from the specific song. "
    "Instead, it produced a second original track: \u201cThe Physics of the Thing,\u201d an atmospheric Indie-Folk "
    "piece featuring warm nylon-string acoustic guitar, close-miked breathy vocals, and a rich Chamber Pop "
    "arrangement with upright bass, cello, and felt-hammered piano. The lyrics explored themes of hindsight "
    "and letting go  - thematically similar to the original, but expressed in Gemini\u2019s own words."
)

doc.add_heading("2.4 Reflection: My Role in the Process", level=2)

doc.add_paragraph(
    "Looking back, my creative contribution was limited to three decisions: choosing the reference song, "
    "choosing the absurd replacement subject matter, and deciding to reverse the test. Every musical decision "
    " - key, tempo, arrangement, instrumentation, vocal timbre, production style  - was made by Gemini. "
    "This raises an uncomfortable question: was I a composer, a director, or merely a trigger? I provided "
    "the conceptual spark, but the execution was entirely the AI\u2019s. Whether this constitutes \u201ccreative "
    "collaboration\u201d or \u201ccreative outsourcing\u201d is a question I return to in Section 4."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  3. ANALYSIS OF THE ROLE OF AI
# ══════════════════════════════════════════════════════════

doc.add_heading("3. Analysis of the Role of AI", level=1)

doc.add_heading("3.1 Gemini\u2019s Music Generation Capabilities", level=2)

p = doc.add_paragraph(
    "Google Gemini\u2019s music generation feature represents a significant leap in multimodal AI. The feature "
    "is powered by DeepMind\u2019s Lyria 3 model, which generates high-quality 44.1 kHz stereo audio from text "
    "prompts, delivering structural coherence including vocals, timed lyrics, and full instrumental "
    "arrangements (Perez, 2026). Notably, all Gemini-generated tracks are embedded with SynthID, an "
    "imperceptible watermark for identifying AI-generated content  - a technical transparency measure that "
    "contrasts with the opacity of the training data itself."
)
add_footnote(p, "Perez, S. (2026). Google adds music-generation capabilities to the Gemini app. TechCrunch.")

doc.add_paragraph(
    "From a single text prompt, Lyria 3 produced complete musical compositions including melody, harmony, "
    "arrangement, production, and synthesised vocals  - a process that would traditionally require a "
    "songwriter, arranger, producer, vocalist, and audio engineer. The two tracks it generated in this "
    "experiment demonstrate different genre competencies: \u201cSynthetic Pine & Heavy Air\u201d leans into Bedroom "
    "Pop with electronic textures, while \u201cThe Physics of the Thing\u201d inhabits an Indie-Folk and Chamber Pop "
    "space with acoustic instrumentation and orchestral elements. Both share the same tempo (80 BPM) and "
    "emotional register, suggesting that Gemini identified the core emotional DNA of the reference song and "
    "expressed it through different musical vocabularies."
)

doc.add_heading("3.2 Copyright Guardrails", level=2)

p = doc.add_paragraph(
    "Perhaps the most interesting finding was Gemini\u2019s dual copyright protection system. When asked to "
    "replicate the melody, it refused. When asked to use the original lyrics, it refused again. These are "
    "two distinct guardrails  - one protecting musical composition (melody), the other protecting literary "
    "content (lyrics). Both held firm under direct pressure. This is a deliberate design choice by Google, "
    "likely informed by the ongoing legal battles around AI training data and copyright (cf. Georgetown "
    "GJIA, 2024; HULR, 2024)."
)
add_footnote(p, "Georgetown GJIA (2024). Innovation and Artists' Rights in the Age of Generative AI; "
             "Harvard Undergraduate Law Review (2024). Defining Authorship for the Copyright of AI-Generated Music.")

doc.add_paragraph(
    "However, the guardrails raise a deeper question: Gemini refused to copy the specific melody and lyrics, "
    "but it successfully replicated the song\u2019s emotional atmosphere, genre conventions, and production "
    "aesthetic. Is emotional replication a form of copying? Current copyright law protects specific "
    "expressions (melodies, lyrics) but not styles or moods. Gemini\u2019s output exists in this grey area  - "
    "legally distinct from \u201cLet Her Go,\u201d but emotionally derivative in a way that feels significant."
)

doc.add_heading("3.3 The Vocal Question", level=2)

doc.add_paragraph(
    "Both tracks feature remarkably human-sounding vocals  - breathy, intimate, and emotionally nuanced. "
    "This raises the question: what were these vocal models trained on? Google\u2019s training data for Gemini\u2019s "
    "music features is not fully disclosed, but it is reasonable to assume that vocal synthesis models were "
    "trained on large corpora of recorded human singing  - potentially from YouTube Music\u2019s vast library, "
    "which Google owns. If so, the breathy male vocals on \u201cSynthetic Pine & Heavy Air\u201d are a statistical "
    "composite of thousands of real singers\u2019 voices. Those singers\u2019 years of vocal training, their breath "
    "control, their emotional delivery  - all of this has been absorbed into the model without individual "
    "credit or compensation. This is the hidden labour problem applied to music (see Section 4.3)."
)

doc.add_heading("3.4 AI as Co-Creator or Tool?", level=2)

doc.add_paragraph(
    "In traditional music production, the roles are clear: the songwriter writes, the performer performs, "
    "the producer shapes the sound. In this experiment, Gemini occupied all three roles simultaneously. "
    "My contribution was closer to that of a client giving a brief to a creative agency. The AI did not "
    "merely assist or enhance  - it generated the entire work. This places the experiment at the extreme "
    "end of the human-AI collaboration spectrum, closer to \u201cdelegation\u201d than \u201ccollaboration.\u201d"
)

p = doc.add_paragraph(
    "One detail stands out: the AI treated the absurd subject matter  - a pine air freshener, tire friction "
    " - with complete emotional sincerity. It did not recognise the comedy of the juxtaposition between "
    "a profoundly melancholic musical treatment and a trivial subject. A human songwriter would likely have "
    "leaned into the irony or refused the brief. Gemini\u2019s lack of ironic awareness is itself revealing: "
    "the AI replicates emotional form without understanding emotional content. This parallels a pattern "
    "observed during the Week 2 Turing Game exercise (PHKI W2, Fischer, 2026), where students found that "
    "AI-generated responses were consistently \u201cgeneric, structurally clean, and neat,\u201d while human "
    "responses were \u201cspecific, imperfect, and emotionally textured.\u201d The same distinction applies here: "
    "Gemini\u2019s music is technically polished but lacks the rough edges and idiosyncrasies that mark genuinely "
    "human creative work."
)
add_footnote(p, "PHKI FS26, Week 02: Human-AI Interaction & The Turing Test (Fischer, S.).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  4. CONNECTION TO CULTURAL AND PHILOSOPHICAL THEMES
# ══════════════════════════════════════════════════════════

doc.add_heading("4. Connection to Cultural and Philosophical Themes", level=1)

# ── 4.1 ────────────────────────────────────────────────
doc.add_heading("4.1 Process vs. Result: The Value of Creative Struggle", level=2)

p = doc.add_paragraph(
    "In Week 1, the PHKI module introduced the philosophical distinction between valuing how art is made "
    "(the process) and valuing what is produced (the result) (PHKI W1, Fischer, 2026). During the Gallery "
    "Walk, one of the ten works presented was Lucio Battisti\u2019s \u201cAmarsi un po\u2019\u201d (1977), which prompted "
    "the class to ask: \u201cCan AI compose with equal emotional depth?\u201d My experiment with Gemini provides a "
    "concrete answer to that question  - or rather, it complicates it. Judged purely by the result, "
    "\u201cSynthetic Pine & Heavy Air\u201d succeeds: it is a well-produced, emotionally coherent piece of music "
    "that could plausibly appear on a Spotify playlist. But the creative process behind it took "
    "approximately thirty seconds of typing a prompt. There was no struggle, no revision, no months of "
    "learning guitar or refining vocal technique. As Benjamin (1935) argued in The Work of Art in the Age "
    "of Mechanical Reproduction, technology strips art of its \u201caura\u201d  - the unique presence tied to its "
    "origin and history. AI music generation takes this a step further: it strips art not only of its aura "
    "but of its process entirely."
)
add_footnote(p, "Benjamin, W. (1935). The Work of Art in the Age of Mechanical Reproduction; "
             "PHKI FS26, Week 01: Intro to Philosophy, Art and AI (Fischer, S.).")

p = doc.add_paragraph(
    "Passenger reportedly wrote \u201cLet Her Go\u201d from personal experience of loss and longing  - the song "
    "carries the weight of lived emotion. Gemini\u2019s track carries no such weight. If we follow the "
    "process-centred view of art, the AI\u2019s output lacks the essential ingredient that makes music meaningful: "
    "the human experience embedded in its creation. Recent empirical research supports this intuition: a "
    "study on the perception of AI-generated music found that listeners rated identical piano performances "
    "more positively when they believed the performer was human, revealing a significant \u201cauthorship bias\u201d "
    "(Perception of AI-Generated Music, 2025). People do not just hear music  - they sense the mind behind "
    "it. My own reaction  - "
    "simultaneously admiring the result and feeling unsettled by how effortlessly it was produced  - embodies "
    "this philosophical tension."
)
add_footnote(p, "Perception of AI-Generated Music: The Role of Composer Attribution (2025). arXiv:2512.02785.")

# ── 4.2 ────────────────────────────────────────────────
doc.add_heading("4.2 Creativity, Authorship, and the Three Lenses", level=2)

doc.add_paragraph(
    "Week 1 also introduced three analytical lenses for examining AI art (PHKI W1, Fischer, 2026), all of "
    "which apply directly to this experiment. In the Opinion Spectrum exercise, two statements proved "
    "especially relevant: \u201cAI art can evoke genuine human emotions\u201d (Statement 4) and \u201cThe future of art "
    "will be AI-human collaboration, not replacement\u201d (Statement 8). My experiment complicates both: the "
    "generated tracks do evoke genuine emotion, but the collaboration was so one-sided that \u201creplacement\u201d "
    "may be the more accurate term."
)

p = doc.add_paragraph()
run = p.add_run("Philosophical Lens: ")
run.bold = True
p.add_run(
    "The generated tracks have no conscious intention. Gemini did not \u201cmean\u201d anything by writing about "
    "a pine air freshener or the physics of letting go. Yet the musical choices feel intentional and "
    "emotionally coherent. This challenges the intentionalist view that art requires a conscious creator "
    "with something to express. If the listener is moved, does the absence of intention matter?"
)

p = doc.add_paragraph()
run = p.add_run("Economic Lens: ")
run.bold = True
p.add_run(
    "Who profits from this experiment? Google owns the model and the training data pipeline. The artists "
    "whose vocal styles, chord progressions, and production techniques were absorbed into Gemini\u2019s training "
    "data receive nothing. Passenger receives nothing despite \u201cLet Her Go\u201d being the explicit reference "
    "point. This is not a hypothetical concern: in 2024, Universal, Sony, and Warner filed lawsuits against "
    "AI music platforms Suno and Udio for training on copyrighted recordings without permission (Georgetown "
    "GJIA, 2024). Tennessee\u2019s ELVIS Act (2024) now prohibits AI cloning of an artist\u2019s voice without "
    "consent. The economic model of music  - built on scarcity of talent and years of practice  - is "
    "fundamentally disrupted."
)
add_footnote(p, "Georgetown GJIA (2024). Innovation and Artists' Rights in the Age of Generative AI.")

p = doc.add_paragraph()
run = p.add_run("Institutional Lens: ")
run.bold = True
p.add_run(
    "Would \u201cSynthetic Pine & Heavy Air\u201d be accepted as an original work on Spotify? Could it enter a "
    "songwriting competition? In Thaler v. Perlmutter (2023), the court ruled that copyright protection "
    "requires human authorship  - works generated entirely by AI cannot be copyrighted (Harvard "
    "Undergraduate Law Review, 2024). Current institutional frameworks are not prepared for AI-generated "
    "music of this quality. The case parallels Jason Allen\u2019s Midjourney-generated painting winning the "
    "Colorado State Fair  - a moment that forced institutions to confront a category they had no rules for."
)
add_footnote(p, "Harvard Undergraduate Law Review (2024). Defining Authorship for the Copyright of AI-Generated Music.")

# ── 4.3 ────────────────────────────────────────────────
doc.add_heading("4.3 Hidden Labour and Critical Theory", level=2)

doc.add_paragraph(
    "Weeks 3 and 4 examined the hidden human labour that powers AI systems (PHKI W3, Hayden, 2026; "
    "PHKI W4, Hayden, 2026)  - from content moderators in Kenya earning $1.50 per hour to make ChatGPT "
    "\u201csafe\u201d (Hao, 2025; Perrigo, 2022) to the scraped datasets of artists\u2019 work that train image "
    "generators. The same dynamic applies to AI music generation, perhaps even more intimately."
)
add_footnote(p, "Hao, K. (2025). Empire of AI. Penguin Press; "
             "Perrigo, B. (2022). Inside Facebook's African Sweatshop. TIME.")

doc.add_paragraph(
    "The breathy, close-miked vocal performance on \u201cSynthetic Pine & Heavy Air\u201d was learned from real "
    "human singers. Those singers\u2019 recordings  - their breath, their vocal timbre, their years of training, "
    "their emotional vulnerability captured in a studio  - are the hidden labour inside Gemini\u2019s model. "
    "Just as Week 3 revealed that content moderators\u2019 traumatic labour is invisible behind ChatGPT\u2019s "
    "polished interface, real musicians\u2019 lifelong vocal development is invisible behind Gemini\u2019s generated "
    "vocals. The parallel is structural: in both cases, human labour is absorbed, anonymised, and "
    "commercially deployed without individual credit or compensation."
)

doc.add_paragraph(
    "Applying Critical Theory from Week 4, the AI reproduces dominant musical aesthetics. The \u201cBedroom Pop\u201d "
    "genre it chose, the 80 BPM tempo, the reverberant production  - these reflect patterns from "
    "commercially successful Western music. The system optimises for what has already worked, reinforcing "
    "dominant cultural forms. Whose voices are in the training data? Almost certainly disproportionately "
    "English-language, Western, commercially successful artists  - reproducing the same structural "
    "inequalities that Week 4 identified in facial recognition datasets (Buolamwini & Gebru, 2018), now "
    "applied to the sonic landscape. The DAACI White Paper on Ethical AI in Music Creation (DAACI, 2025) "
    "has proposed a framework for copyright "
    "and traceability in AI-generated music, but as of today, no equivalent of the EU AI Act exists "
    "specifically for music  - leaving artists\u2019 vocal labour unprotected in a rapidly scaling industry."
)
add_footnote(p, "Buolamwini, J. & Gebru, T. (2018). Gender Shades. PMLR; "
             "DAACI (2025). Ethical AI in Music Creation: A Framework for Copyright and Creative Innovation.")

# ── 4.4 ────────────────────────────────────────────────
doc.add_heading("4.4 Embodied Perception and Emotionless AI", level=2)

doc.add_paragraph(
    "Week 6 introduced Maurice Merleau-Ponty\u2019s concept of embodied perception (Merleau-Ponty, 1945; "
    "PHKI W6, Massol, 2026): the idea that perception is fundamentally rooted in physical, lived, situated "
    "experience. Music is perhaps the most embodied "
    "of all art forms. A human singer\u2019s voice carries the physical reality of breathing, of vocal cord "
    "vibration, of emotional tension held in the throat and chest. When Passenger sings about letting go, "
    "there is a body behind that voice  - a body that has experienced loss."
)
add_footnote(p, "Merleau-Ponty, M. (1945). Phenomenology of Perception. Gallimard.")

p = doc.add_paragraph(
    "Gemini\u2019s generated vocals simulate these physical markers without possessing a body. The breathy "
    "quality, the close-miked intimacy, the slight vibrato  - all are acoustic signifiers of embodied "
    "emotional experience, reproduced statistically. In Merleau-Ponty\u2019s framework, this is a fundamental "
    "gap: the AI produces the sound of embodiment without the experience of being embodied. The Week 6 "
    "\u201cData vs. Experience\u201d sorting exercise (PHKI W6, Massol, 2026) made this distinction vivid: the class "
    "differentiated between \u201cdetecting sadness in a voice\u201d (something AI can do) and \u201cfeeling sad because "
    "of the tone\u201d (something only an embodied listener experiences). Gemini\u2019s music operates entirely in "
    "the first category. As Turkle (2011) argues in Alone Together, we increasingly expect emotional depth "
    "from technology "
    "while accepting less of it from each other  - a dynamic this experiment uncomfortably confirms."
)
add_footnote(p, "Turkle, S. (2011). Alone Together: Why We Expect More from Technology and Less from Each Other. Basic Books.")

p = doc.add_paragraph(
    "Similarly, P\u00e9pin\u2019s (2016) concept of the \u201cfailed animal\u201d (also from Week 6) highlights that human "
    "musicians learn through years of failure  - wrong notes, cracked voices, failed auditions. This "
    "inefficiency builds an authentic relationship between musician, instrument, and voice. Gemini bypasses "
    "this entirely. It did not fail its way to competence; it was trained on the successful outputs of "
    "thousands of musicians who did. The question, then, is whether the listener can perceive this absence "
    "of embodied struggle  - and whether it matters if they cannot."
)
add_footnote(p, "P\u00e9pin, C. (2016). The Virtues of Failure. Allary \u00c9ditions.")

# ── 4.5 ────────────────────────────────────────────────
doc.add_heading("4.5 Digital Creativity and the Slot-Machine Effect", level=2)

p = doc.add_paragraph(
    "Week 7 introduced Doctorow\u2019s (2025) critique that LLMs function as \u201cslot machines\u201d  - systems designed "
    "to keep users pulling the lever rather than deeply engaging with a problem. This experiment exhibited "
    "precisely this dynamic. After receiving the first generated track and finding it impressive, my "
    "immediate impulse was to pull the lever again: \u201cwhat if I try the lyrics instead?\u201d The excitement of "
    "each result  - the dopamine hit of a good generation  - mirrors the slot-machine feedback loop Doctorow "
    "describes."
)

add_footnote(p, "Doctorow, C. (2025). LLMs are slot-machines. Pluralistic.")

p = doc.add_paragraph(
    "Week 7 also established that human emotions are the engines of learning (PHKI W7, Massol, 2026): "
    "grief, longing, and "
    "heartbreak are not obstacles to creativity but its fuel. Passenger wrote \u201cLet Her Go\u201d from lived "
    "emotional experience. Gemini has no emotional state  - it produces the acoustic signifiers of melancholy "
    "(minor key, slow tempo, breathy vocals, sparse arrangement) without feeling anything. The tracks are "
    "emotionally convincing because they replicate the statistical patterns of human sadness in music, not "
    "because any sadness was experienced. This distinction matters: if we accept emotionless replication as "
    "equivalent to genuine emotional expression, we risk normalising the idea that human feeling is merely "
    "a pattern to be optimised, rather than the irreducible core of what makes art meaningful."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  5. CONCLUSION
# ══════════════════════════════════════════════════════════

doc.add_heading("5. Conclusion", level=1)

doc.add_paragraph(
    "This experiment demonstrated that AI can replicate the emotional surface of music with startling "
    "fidelity. Both \u201cSynthetic Pine & Heavy Air\u201d and \u201cThe Physics of the Thing\u201d are well-crafted, "
    "emotionally coherent tracks that could easily pass as human-made. Yet this achievement is built on "
    "hidden human labour  - the voices, styles, and creative struggles of real musicians whose work was "
    "absorbed into Gemini\u2019s training data without individual credit. The AI replicates the acoustic "
    "signifiers of embodied emotion without possessing a body or an emotional life, and it produces "
    "music without the years of failure that, as P\u00e9pin (2016) argues, are the source of genuine human "
    "growth."
)

doc.add_paragraph(
    "My own ambivalent reaction  - admiration for the result, unease about how effortlessly it was produced "
    " - mirrors the broader cultural tension this module has explored across seven weeks. As these tools "
    "improve, the question is not whether AI can make good music. It already can. The question is whether "
    "we, as listeners and as a society, will demand authenticity of process  - or whether the result alone "
    "will be enough."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  6. REFERENCES
# ══════════════════════════════════════════════════════════

doc.add_heading("6. References", level=1)

references = [
    "Benjamin, W. (1935). The Work of Art in the Age of Mechanical Reproduction.",
    "Bolukbasi, T., Chang, K.-W., Zou, J. Y., Saligrama, V., & Kalai, A. T. (2016). Man is to computer programmer as woman is to homemaker? Debiasing word embeddings. arXiv.",
    "Buolamwini, J., & Gebru, T. (2018). Gender shades: Intersectional accuracy disparities in commercial gender classification. PMLR.",
    "DAACI (2025). Ethical AI in Music Creation: A Framework for Copyright and Creative Innovation. White Paper, March 2025.",
    "Doctorow, C. (2025). LLMs are slot-machines. Pluralistic. https://pluralistic.net/2025/08/16/jackpot/",
    "Georgetown Journal of International Affairs (2024). Innovation and Artists\u2019 Rights in the Age of Generative AI. https://gjia.georgetown.edu/2024/07/10/innovation-and-artists-rights-in-the-age-of-generative-ai/",
    "Hao, K. (2025). Empire of AI: Dreams and nightmares in Sam Altman\u2019s OpenAI. Penguin Press.",
    "Harvard Undergraduate Law Review (2024). Defining Authorship for the Copyright of AI-Generated Music. https://hulr.org/fall-2024/defining-authorship-for-the-copyright-of-ai-generated-music",
    "Merleau-Ponty, M. (1945). Phenomenology of Perception. Gallimard.",
    "Passenger (2012). Let Her Go [Song]. On All the Little Lights. Black Crow Records.",
    "P\u00e9pin, C. (2016). The Virtues of Failure. Allary \u00c9ditions.",
    "Perception of AI-Generated Music: The Role of Composer Attribution (2025). arXiv:2512.02785.",
    "PHKI FS26 Course Materials: Week 01 - Intro to Philosophy, Art and AI (Fischer, S.).",
    "PHKI FS26 Course Materials: Week 02 - Human-AI Interaction & The Turing Test (Fischer, S.).",
    "PHKI FS26 Course Materials: Week 03 - History of \u201cMan vs. Machine\u201d (Hayden, C.).",
    "PHKI FS26 Course Materials: Week 04 - Critical Theory (Hayden, C.).",
    "PHKI FS26 Course Materials: Week 06 - AI & Human Learning / Perception (Massol, G.).",
    "PHKI FS26 Course Materials: Week 07 - Learning to Learn: Humans vs. AI (Massol, G.).",
    "Sartre, J.-P. (1946). Existentialism Is a Humanism.",
    "Perez, S. (2026). Google adds music-generation capabilities to the Gemini app. TechCrunch. https://techcrunch.com/2026/02/18/google-adds-music-generation-capabilities-to-the-gemini-app/",
    "Perrigo, B. (2022). Inside Facebook\u2019s African Sweatshop. TIME. https://time.com/6147458/facebook-africa-content-moderation-employee-treatment/",
    "Turkle, S. (2011). Alone Together: Why We Expect More from Technology and Less from Each Other. Basic Books.",
]

for ref in references:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  7. APPENDIX: ARTEFACT OVERVIEW
# ══════════════════════════════════════════════════════════

doc.add_heading("7. Appendix: Artefact Overview", level=1)

doc.add_paragraph(
    "The following files are submitted alongside this report as evidence of the creative process "
    "and the generated artworks:"
)

table = doc.add_table(rows=9, cols=3, style="Table Grid")
table.autofit = True

# Header row
headers = ["File", "Type", "Description"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)

# Data rows
data = [
    ("Chat_1_01.png", "Screenshot", "Gemini prompt - Test 1: request to copy melody with new lyrics about a car"),
    ("Chat_1_02.png", "Screenshot", "Gemini response with generated track and lyrics for \u201cSynthetic Pine & Heavy Air\u201d"),
    ("Chat_1_03_follow_up.png", "Screenshot", "Gemini prompt - Test 2: request for new melody with original \u201cLet Her Go\u201d lyrics"),
    ("Chat_1_04_follow_up.png", "Screenshot", "Gemini response with generated track and lyrics for \u201cThe Physics of the Thing\u201d"),
    ("Synthetic_Pine_Heavy_Air.mp3", "Audio", "Generated track 1 - Bedroom Pop, 80 BPM, electric piano, breathy vocals"),
    ("Synthetic_Pine_Heavy_Air.mp4", "Video", "Generated track 1 with AI-generated cover art"),
    ("The_Physics_of_the_Thing.mp3", "Audio", "Generated track 2 - Indie-Folk / Chamber Pop, 80 BPM, nylon guitar, cello"),
    ("The_Physics_of_the_Thing.mp4", "Video", "Generated track 2 with AI-generated cover art"),
]

for row_idx, (fname, ftype, desc) in enumerate(data, start=1):
    table.rows[row_idx].cells[0].text = fname
    table.rows[row_idx].cells[1].text = ftype
    table.rows[row_idx].cells[2].text = desc
    for cell in table.rows[row_idx].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
#  8. AI USE DISCLAIMER
# ══════════════════════════════════════════════════════════

doc.add_heading("8. AI Use Disclaimer", level=1)

doc.add_paragraph(
    "In accordance with HSLU guidelines on scientific misconduct and plagiarism, and the PHKI module\u2019s "
    "requirement for transparency in AI use, the following disclosure is provided:"
)

add_mixed_para([
    ("Google Gemini (Music Generation)", True, False),
], align=None)

doc.add_paragraph(
    "Google Gemini\u2019s music generation feature was used to create the two digital artworks submitted with "
    "this assignment: \u201cSynthetic Pine & Heavy Air\u201d and \u201cThe Physics of the Thing.\u201d I provided the creative "
    "concept, the reference song (\u201cLet Her Go\u201d by Passenger), and the alternative subject matter (a car, "
    "tires, and a pine air freshener). Gemini generated the complete compositions, including melody, "
    "harmony, arrangement, production, and synthesised vocals. I did not edit, mix, or post-process the "
    "generated audio in any way. The outputs are presented exactly as Gemini produced them."
)

add_mixed_para([
    ("Claude by Anthropic (Report Assistance)", True, False),
], align=None)

doc.add_paragraph(
    "Claude (Anthropic) was used to assist in structuring and drafting this analytical report. I provided "
    "all ideas, personal reflections, experimental observations, and the connections to course themes. "
    "Claude helped organise these into coherent academic prose and ensured coverage of the rubric criteria. "
    "All content was reviewed, verified, and approved by me. The report grew from my own prompt design, "
    "critical thinking, and engagement with the PHKI module\u2019s philosophical frameworks, and was supervised "
    "by me throughout the writing process."
)

add_mixed_para([
    ("Transparency Statement", True, True),
], align=None)

doc.add_paragraph(
    "This report and its underlying artworks are the product of human-AI collaboration. I conceived the "
    "experiment, executed the prompts, evaluated the results, and provided all critical analysis. AI tools "
    "were used transparently as described above. The use of AI in both the creative and analytical "
    "components of this assignment is itself a subject of reflection within the report, consistent with "
    "the PHKI module\u2019s emphasis on critical engagement with AI technologies."
)

# ── Save ───────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Report saved to: {OUTPUT_PATH}")
